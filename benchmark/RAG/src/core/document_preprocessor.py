"""Format-neutral document preparation before OpenViking ingestion."""

from __future__ import annotations

import hashlib
import os
import re
from pathlib import Path


class DocumentPreprocessor:
    """Convert common local document formats to stable Markdown text.

    Plain text and Markdown are passed through unchanged.  Conversion is
    selected only by file format, never by benchmark or dataset name.
    """

    _PASSTHROUGH_SUFFIXES = {".md", ".markdown", ".txt", ".text", ".rst"}

    def prepare(self, path: str, output_dir: str | None = None) -> str:
        source = Path(path).expanduser().resolve()
        suffix = source.suffix.lower()
        if suffix in self._PASSTHROUGH_SUFFIXES:
            return str(source)
        if suffix not in {".pdf", ".docx"}:
            return str(source)

        target_dir = Path(output_dir).expanduser() if output_dir else source.parent / ".openviking_preprocessed"
        target_dir.mkdir(parents=True, exist_ok=True)
        digest = hashlib.sha1(str(source).encode("utf-8")).hexdigest()[:10]
        target = target_dir / f"{source.stem}_{digest}.md"
        if target.exists() and target.stat().st_mtime >= source.stat().st_mtime:
            return str(target)

        if suffix == ".pdf":
            text = self._pdf_to_markdown(source)
        else:
            text = self._docx_to_markdown(source)
        target.write_text(text, encoding="utf-8")
        return str(target)

    @staticmethod
    def _clean_text(text: str) -> str:
        text = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip() + "\n"

    def _pdf_to_markdown(self, source: Path) -> str:
        try:
            import pdfplumber
        except ImportError as exc:
            raise RuntimeError(
                "PDF ingestion requires pdfplumber; install it in the benchmark environment"
            ) from exc

        pages = []
        with pdfplumber.open(str(source)) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(f"## Page {page_number}\n\n{page_text}")
        return self._clean_text("\n\n".join(pages))

    def _docx_to_markdown(self, source: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "DOCX ingestion requires python-docx; install it in the benchmark environment"
            ) from exc

        document = Document(str(source))
        blocks = []
        for paragraph in document.paragraphs:
            value = paragraph.text.strip()
            if value:
                blocks.append(value)
        for table in document.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells))
            if rows:
                blocks.append("\n".join(rows))
        return self._clean_text("\n\n".join(blocks))
